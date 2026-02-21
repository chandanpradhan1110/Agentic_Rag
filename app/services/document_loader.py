"""
app/services/document_loader.py

IMPLEMENT HERE:
- load_document()  → reads PDF / TXT / DOCX / CSV and returns raw text
- chunk_text()     → splits text into overlapping word-level chunks
"""
import re
from pathlib import Path


def load_document(file_path: Path, filename: str) -> str:
    """
    Dispatch to the right loader based on file extension.
    Returns the full extracted text string.
    """
    ext = file_path.suffix.lower()
    loaders = {
        ".pdf":  _load_pdf,
        ".txt":  _load_text,
        ".md":   _load_text,
        ".docx": _load_docx,
        ".csv":  _load_csv,
    }
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file type: {ext}")

    text = loader(file_path)
    return _clean(text)


# ── Loaders ───────────────────────────────────────────────────────────────────

def _load_pdf(path: Path) -> str:
    """
    IMPLEMENT: Use pypdf.PdfReader to extract text page by page.
    Return "\n\n".join of all page texts.
    """
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text()
            if t and t.strip():
                pages.append(f"[Page {i+1}]\n{t.strip()}")
        except Exception:
            pass
    return "\n\n".join(pages)


def _load_text(path: Path) -> str:
    """
    IMPLEMENT: Try UTF-8, fall back to latin-1.
    """
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _load_docx(path: Path) -> str:
    """
    IMPLEMENT: Use python-docx. Extract paragraph text.
    Prefix Heading paragraphs with ## for context.
    Also extract tables as pipe-delimited rows.
    """
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if para.style.name.startswith("Heading"):
            parts.append(f"\n## {t}\n")
        else:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n\n".join(parts)


def _load_csv(path: Path) -> str:
    """
    IMPLEMENT: Use pandas. Convert each row to a key=value sentence.
    Cap at 1000 rows for large files.
    """
    import pandas as pd
    df = pd.read_csv(str(path))
    lines = [f"Columns: {', '.join(df.columns)}"]
    for _, row in df.head(1000).iterrows():
        lines.append("; ".join(f"{k}={v}" for k, v in row.items()))
    if len(df) > 1000:
        lines.append(f"... {len(df)-1000} more rows truncated")
    return "\n".join(lines)


# ── Chunker ───────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 512, overlap: int = 64) -> list[str]:
    """
    IMPLEMENT: Sentence-aware sliding-window chunker.

    Algorithm:
    1. Split text into sentences on [.!?] boundaries
    2. Accumulate sentences into a chunk until chunk_size words exceeded
    3. When flushing, keep last `overlap` words as the start of next chunk
    4. Filter out chunks with fewer than 5 words

    Returns list of chunk strings.
    """
    if not text or not text.strip():
        return []

    sentences = re.split(r"(?<=[.!?])\s+", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for sentence in sentences:
        words = sentence.split()
        if current_len + len(words) > chunk_size and current:
            chunks.append(" ".join(current))
            current = current[-overlap:] + words
            current_len = len(current)
        else:
            current.extend(words)
            current_len += len(words)

    if current:
        chunks.append(" ".join(current))

    return [c for c in chunks if len(c.split()) >= 5]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r" {3,}", "  ", text)
    return text.strip()
